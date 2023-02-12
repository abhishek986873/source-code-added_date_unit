# mport statement
import docx as docx

# ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# FIXING PRICES VARIABLE


dmls = 600
metal_free = 1800
pfm = 350
flexible = 0
rpd = 600
ortho = 700
repair = 150
dental_repair = 150
cd = 600
zirconia = 1800

# ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------


# create a new document
mintu_bill = docx.Document()

# ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

# add a table to the document with one row and one column FOR HEADING
table = mintu_bill.add_table(rows=1, cols=1)

# ----------------------- ADDING - HEADING -------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

cell = table.cell(0, 0)

# Add the main heading
main_heading = cell.add_paragraph("ORO CARE DENTAL LAB\n H.NO- 223, GALI NO.-2, SAIDULJAB, SAKET, NEW DELHI- 110030")
main_heading.style = "Title"

# Add the subheading
sub_heading = cell.add_paragraph(
    "Ref.........                                                                              Date-31-01-2023 \n\nINSTITUTE OF PARAMEDICAL TECHNOLOGY \nCHATTERPUR , NEW DELHI\n")
sub_heading.style = "Subtitle"

# --------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

# TAKING INPUT OF NUMBER OF ROW IN A TABLE

bill_row = int(input("ENTER NUMBER OF ROW = "))
print(
    "------------------------------------------------⚫ UNIT----------------------------------------------------------------------------------------------------")

# ADDING 1 TO ROW VALUE BECAUSE WE HAVE TO CONSIDER THE OTH INDEX THAT IS CONSIST OF HEADING LIKE UNIT , S.NO , DATE

bill_row += 1
table = mintu_bill.add_table(rows=bill_row, cols=5)

# ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

# ADDING HEADING IN THE FIRST ROW OF MINTU TABLE

heading_row = table.rows[0].cells

heading_row[0].text = "S.NO"
heading_row[1].text = "DATE"
heading_row[2].text = "UNIT"
heading_row[3].text = "WORK"
heading_row[4].text = "AMOUNT"

# ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

# filling the SERIAL NUMBER COLUMN

i = 1
k = 1
while i < bill_row:
    data_row = table.rows[i].cells
    data_row[0].text = str(k)
    k = k + 1
    i = i + 1

# ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------


# Creating list for greeting the unit

unit_value_list = ["ENTER 1ST UNIT = ", "ENTER 2ND UNIT = ", "ENTER 3RD UNIT = "]
i = 4
while i < bill_row:
    string = f"ENTER {i}TH UNIT = "
    unit_value_list.append(string)
    i = i + 1

# ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

#  entering the value of unit inside the cell


# creating empty list to store value of unit


unit_list = []
i = 1
j = 0
while i < bill_row:
    data_row = table.rows[i].cells
    unit_value = data_row[2].text = input(unit_value_list[j])
    unit_list.append(unit_value)
    i = i + 1
    j = j + 1


# ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

print(
    "------------------------------------------------⚫ DATE----------------------------------------------------------------------------------------------------")

# TAKING INPUT OF MONTH FOR DATE COLOUMN

month = input("ENTER MONTH = ")

print(
    "--------------------------------------------------------------------------------------------------------------------------------------------------")

# Creating list for greeting the DATE

date_value_list = ["ENTER 1ST DATE = ", "ENTER 2ND DATE = ", "ENTER 3RD DATE = "]
i = 4
while i < bill_row:
    string_date = f"ENTER {i}TH DATE = "
    date_value_list.append(string_date)
    i = i + 1

#  --------------------------------inserting date in a cell--------------------------------------------------------------------------
i = 1
j = 0
while i < bill_row:
    data_row = table.rows[i].cells
    data_row[1].text = str(input(date_value_list[j]) + f"-{month}-2023")
    i = i + 1
    j = j + 1

# ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

print(
    "------------------------------------------------⚫ WORK----------------------------------------------------------------------------------------------------")

# creating greeting for work

work_value_list = ["ENTER 1ST WORK = ", "ENTER 2ND WORK = ", "ENTER 3RD WORK = "]
i = 4
while i < bill_row:
    string_work = f"ENTER {i}TH WORK = "
    work_value_list.append(string_work)
    i = i + 1

# inserting value inside work cell

# inserting the work inside a list

work_list = []
i = 1
j = 0
while i < bill_row:
    data_row = table.rows[i].cells
    work_value = data_row[3].text = input(work_value_list[j])
    work_list.append(work_value)
    i = i + 1
    j = j + 1

# --------------------------------------AMOUNT--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# creating list of amount value

amount_value_list = []
i = 0
j = 1
while j < bill_row:
    if work_list[i] == "PFM":
        amount_value = int(unit_list[i]) * pfm
        amount_value_list .append(amount_value)

    elif work_list[i] == "DMLS":
        amount_value = int(unit_list[i]) * dmls
        amount_value_list .append(amount_value)

    elif work_list[i] == "METAL FREE":
        amount_value = int(unit_list[i]) * metal_free
        amount_value_list .append(amount_value)

    elif work_list[i] == "FLEXIBLE":
        amount_value = int(unit_list[i]) * flexible
        amount_value_list .append(amount_value)

    elif work_list[i] == "RPD":
        amount_value = int(unit_list[i]) * rpd
        amount_value_list .append(amount_value)

    elif work_list[i] == "ORTHO":
        amount_value = int(unit_list[i]) * ortho
        amount_value_list .append(amount_value)

    elif work_list[i] == "REPAIR":
        amount_value = int(unit_list[i]) * repair
        amount_value_list .append(amount_value)

    elif work_list[i] == "DENTAL REPAIR":
        amount_value = int(unit_list[i]) * dental_repair
        amount_value_list .append(amount_value)

    elif work_list[i] == "CD":
        amount_value = int(unit_list[i]) * cd
        amount_value_list .append(amount_value)

    elif work_list[i] == "ZIRCONIA":
        amount_value = int(unit_list[i]) * zirconia
        amount_value_list .append(amount_value)
    i = i + 1
    j = j + 1

# ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# inserting amount value


i = 1
j = 0
while i < bill_row:
    data_row = table.rows[i].cells
    data_row[4].text = str(amount_value_list[j])

    i = i + 1
    j = j + 1

# ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

# CALCULATING TOTAL AMOUNT

total_amount = sum(amount_value_list)

# ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# entering the value of total and greeting inside mintu table











# ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

table = mintu_bill.add_table(rows=2, cols=5)
data_row = table.rows[1].cells
data_row[3].text = "TOTAL"
data_row = table.rows[1].cells
data_row[4].text = str(total_amount)
mintu_bill.save("mintuttt_bill.docx")


# data_row = table.rows[2].cells
# data_row[4].text = "TOTAL AMOUNT"
